import json
from collections import defaultdict
from datetime import date
from pathlib import Path
from urllib.parse import urlparse

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Project_Sekai_全服务器_MV链接与创作者汇总.docx"

SERVERS = {
    "JP": ROOT / "sekai-master-db-diff-main",
    "CN": ROOT / "sekai-master-db-cn-diff-main",
}

ALLOWED_HOSTS = {
    "youtube.com", "www.youtube.com", "youtu.be", "m.youtube.com",
    "nicovideo.jp", "www.nicovideo.jp", "nico.ms",
    "bilibili.com", "www.bilibili.com", "b23.tv",
}


def load_json(path):
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def clean_url(value):
    if not isinstance(value, str):
        return None
    value = value.strip()
    if not value:
        return None
    host = urlparse(value).netloc.lower()
    return value if host in ALLOWED_HOSTS else None


def unique(values):
    seen = set()
    result = []
    for value in values:
        if value and value not in seen:
            seen.add(value)
            result.append(value)
    return result


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relation_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.append(color)
    run_props.append(underline)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(run_props)
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_repeat_table_header(row):
    tr_props = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_props.append(header)


def shade_cell(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    shade = props.find(qn("w:shd"))
    if shade is None:
        shade = OxmlElement("w:shd")
        props.append(shade)
    shade.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    props = cell._tc.get_or_add_tcPr()
    margins = props.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        props.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    props = table._tbl.tblPr
    width = props.find(qn("w:tblW"))
    width.set(qn("w:w"), str(sum(widths_dxa)))
    width.set(qn("w:type"), "dxa")
    indent = props.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        props.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(value))
        grid.append(col)
    for row in table.rows:
        for cell, value in zip(row.cells, widths_dxa):
            cell.width = Inches(value / 1440)
            tc_width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_width.set(qn("w:w"), str(value))
            tc_width.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def creator_payload(music, artists):
    info = music.get("infos") or []
    info = info[0] if info and isinstance(info[0], dict) else {}
    artist = artists.get(music.get("creatorArtistId"), "")
    return {
        "artist": info.get("creator") or artist or "-",
        "lyricist": info.get("lyricist") or music.get("lyricist") or "-",
        "composer": info.get("composer") or music.get("composer") or "-",
        "arranger": info.get("arranger") or music.get("arranger") or "-",
    }


def format_creator(payload):
    return (
        f"创作者：{payload['artist']}\n"
        f"作词：{payload['lyricist']}\n"
        f"作曲：{payload['composer']}\n"
        f"编曲：{payload['arranger']}"
    )


def collect_rows():
    server_music = {}
    originals_by_server = {}
    for server, folder in SERVERS.items():
        musics = load_json(folder / "musics.json")
        artists = {item["id"]: item.get("name", "") for item in load_json(folder / "musicArtists.json")}
        originals = load_json(folder / "musicOriginals.json")
        original_map = defaultdict(list)
        for item in originals:
            url = clean_url(item.get("videoLink"))
            if url:
                original_map[item.get("musicId")].append(url)
        server_music[server] = {
            item["id"]: {"raw": item, "creator": creator_payload(item, artists)} for item in musics
        }
        originals_by_server[server] = original_map

    database = load_json(ROOT / "output" / "database_v2.json")
    video_by_music_id = defaultdict(lambda: {"original": [], "2dmv": []})
    video_by_title = defaultdict(lambda: {"original": [], "2dmv": []})
    for song in database.get("songs", []):
        destinations = []
        if song.get("sekaiMusicId") is not None:
            destinations.append(video_by_music_id[song["sekaiMusicId"]])
        destinations.append(video_by_title[song.get("titleJp") or song.get("title") or ""])
        for video in song.get("videos", []):
            url = clean_url(video.get("url"))
            if not url:
                continue
            kind = "2dmv" if video.get("type") == "official_2dmv" else "original" if video.get("type") == "original_mv" else None
            if kind:
                for destination in destinations:
                    destination[kind].append(url)

    rows = []
    all_ids = sorted(set(server_music["JP"]) | set(server_music["CN"]))
    for music_id in all_ids:
        present = [server for server in SERVERS if music_id in server_music[server]]
        jp = server_music["JP"].get(music_id)
        cn = server_music["CN"].get(music_id)
        reference = jp or cn
        title = reference["raw"].get("title", "")

        original_links = []
        for server in present:
            original_links.extend(originals_by_server[server].get(music_id, []))
        original_links.extend(video_by_music_id[music_id]["original"])
        original_links.extend(video_by_title[title]["original"])
        mv_links = video_by_music_id[music_id]["2dmv"] + video_by_title[title]["2dmv"]
        original_links = unique(original_links)
        mv_links = unique(mv_links)
        if not original_links and not mv_links:
            continue

        creator_lines = []
        if jp:
            creator_lines.append(("JP", jp["creator"]))
        if cn and (not jp or cn["creator"] != jp["creator"]):
            creator_lines.append(("CN", cn["creator"]))
        creators = "\n".join(
            ((f"[{server}]\n" if len(creator_lines) > 1 else "") + format_creator(payload))
            for server, payload in creator_lines
        )
        rows.append({
            "servers": "/".join(present),
            "music_id": music_id,
            "title": title,
            "original": original_links,
            "mv2d": mv_links,
            "creators": creators,
        })
    return rows, {server: len(items) for server, items in server_music.items()}


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, before, after, color in (
        ("Title", 24, 0, 8, "17365D"),
        ("Heading 1", 16, 14, 8, "2E74B5"),
        ("Heading 2", 13, 10, 5, "1F4D78"),
    ):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])
    paragraph.add_run(" 页")


def build_document(rows, server_counts):
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    header = section.header.paragraphs[0]
    header.text = "Project Sekai MV 链接与游戏内创作者资料"
    header.style = doc.styles["Normal"]
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    add_page_number(section.footer.paragraphs[0])

    title = doc.add_paragraph(style="Title")
    title.add_run("Project Sekai 全服务器 MV 链接与创作者汇总")
    subtitle = doc.add_paragraph()
    subtitle.add_run(f"数据日期：{date.today().isoformat()}  |  收录服务器：JP（日服）、CN（国服）").bold = True
    doc.add_paragraph(
        f"共收录 {len(rows)} 首含有效 MV 链接的歌曲。服务器曲库规模：JP {server_counts['JP']} 首，CN {server_counts['CN']} 首。"
    )
    note = doc.add_paragraph()
    note.add_run("口径说明：").bold = True
    note.add_run(
        "仅保留 YouTube、niconico 与 bilibili 链接；原曲 MV 合并游戏 Master 的 musicOriginals 与本家 MV 数据，"
        "2DMV 取项目官方 2DMV 数据。CN Master 当前没有独立 musicOriginals 链接，因此国服已上线的同 ID 歌曲沿用同一作品的公共 MV 链接。"
    )

    doc.add_paragraph("歌曲明细", style="Heading 1")
    headers = ["服务器", "ID", "歌名", "原曲 MV", "2DMV", "游戏数据内创作者信息"]
    widths = [700, 560, 1550, 2650, 2650, 2290]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_repeat_table_header(table.rows[0])
    for cell, label in zip(table.rows[0].cells, headers):
        shade_cell(cell, "E8EEF5")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(8.5)

    for item in rows:
        cells = table.add_row().cells
        values = [item["servers"], str(item["music_id"]), item["title"]]
        for index, value in enumerate(values):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if index < 2 else WD_ALIGN_PARAGRAPH.LEFT
            p.add_run(value)
        for index, links in ((3, item["original"]), (4, item["mv2d"])):
            cell = cells[index]
            cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
            if links:
                for link_index, url in enumerate(links, 1):
                    p = cell.add_paragraph()
                    p.paragraph_format.space_after = Pt(1)
                    label = f"链接 {link_index}" if len(links) > 1 else "打开链接"
                    add_hyperlink(p, label, url)
                    p.add_run(f"\n{url}")
            else:
                cell.add_paragraph("-")
        creator_cell = cells[5]
        creator_cell.text = item["creators"]
        for cell in cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(7.5)

    set_table_geometry(table, widths)
    doc.core_properties.title = "Project Sekai 全服务器 MV 链接与创作者汇总"
    doc.core_properties.subject = "原曲 MV、2DMV 与游戏数据创作者信息"
    doc.core_properties.author = "Project Sekai 2DMV Database"
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    data_rows, counts = collect_rows()
    output = build_document(data_rows, counts)
    print(json.dumps({"output": str(output), "rows": len(data_rows), "servers": counts}, ensure_ascii=False))
